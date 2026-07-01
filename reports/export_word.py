"""
导出 Word 计算书
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Microsoft YaHei')


def _add_kv_table(doc, data, col_widths=None):
    table = doc.add_table(rows=len(data), cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(data.items()):
        c0 = table.cell(i, 0)
        c0.text = str(k)
        for p in c0.paragraphs:
            for run in p.runs:
                run.bold = True
        c1 = table.cell(i, 1)
        c1.text = str(v)
    return table


def _add_data_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = str(h)
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            table.cell(i + 1, j).text = str(v)
    return table


def export_to_word(collected: dict, filepath: str):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)
    title = doc.add_heading("\u7ba1\u5f0f\u819c\u8bbe\u8ba1\u8ba1\u7b97\u4e66", level=0)
    for run in title.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Microsoft YaHei')
    doc.add_paragraph()
    # project
    if "\u9879\u76ee\u57fa\u7840\u4fe1\u606f" in collected:
        _add_heading(doc, "\u4e00\u3001\u9879\u76ee\u57fa\u7840\u4fe1\u606f", level=1)
        pd = collected["\u9879\u76ee\u57fa\u7840\u4fe1\u606f"]
        info = {
            "\u9879\u76ee\u540d\u79f0": pd.get("project_name", ""),
            "\u9879\u76ee\u7f16\u53f7": pd.get("project_no", ""),
            "\u8bbe\u8ba1\u4ea7\u6c34\u91cf (m\u00b3/d)": pd.get("design_flow", ""),
            "\u6c34\u578b\u5206\u7c7b": pd.get("water_type", ""),
            "\u8bbe\u8ba1\u4eba\u5458": pd.get("designer", ""),
            "\u8bbe\u8ba1\u65e5\u671f": pd.get("design_date", ""),
        }
        _add_kv_table(doc, info)
        doc.add_paragraph()
    # membrane
    if "\u819c\u578b\u53f7" in collected:
        _add_heading(doc, "\u4e8c\u3001\u819c\u578b\u53f7\u53c2\u6570", level=1)
        _add_kv_table(doc, collected["\u819c\u578b\u53f7"])
        doc.add_paragraph()
    # area
    if "\u819c\u9762\u79ef\u8ba1\u7b97" in collected:
        _add_heading(doc, "\u4e09\u3001\u819c\u9762\u79ef\u8ba1\u7b97", level=1)
        _add_kv_table(doc, collected["\u819c\u9762\u79ef\u8ba1\u7b97"])
        doc.add_paragraph()
    # crossflow
    if "\u9519\u6d41\u5faa\u73af\u91cf" in collected:
        _add_heading(doc, "\u56db\u3001\u9519\u6d41\u5faa\u73af\u91cf\u8ba1\u7b97", level=1)
        _add_kv_table(doc, collected["\u9519\u6d41\u5faa\u73af\u91cf"])
        doc.add_paragraph()
    # TMP
    if "TMP\u538b\u529b\u6821\u6838" in collected:
        _add_heading(doc, "\u4e94\u3001TMP\u538b\u529b\u6821\u6838", level=1)
        rd = collected["TMP\u538b\u529b\u6821\u6838"]
        tmp_data = {k: ("; ".join(v) if k == "warnings" else v) for k, v in rd.items()}
        _add_kv_table(doc, tmp_data)
        doc.add_paragraph()
    # pumps
    if "\u6cf5\u9009\u578b" in collected:
        _add_heading(doc, "\u516d\u3001\u6cf5\u9009\u578b", level=1)
        rd = collected["\u6cf5\u9009\u578b"]
        headers = ["\u6cf5\u540d\u79f0", "\u6d41\u91cf(m\u00b3/h)", "\u626c\u7a0b(m)", "\u8f74\u529f\u7387(kW)", "\u7535\u673a\u529f\u7387(kW)", "\u6750\u8d28\u5efa\u8bae"]
        rows = []
        for key in ["feed_pump", "crossflow_pump", "cleaning_pump"]:
            p = rd.get(key, {})
            rows.append([p.get("name", key), f"{p.get('flow', 0):.1f}", f"{p.get('head', 0):.0f}", f"{p.get('shaft_power_kw', 0):.2f}", f"{p.get('recommended_motor_kw', 0):.0f}", p.get("material", "")])
        _add_data_table(doc, headers, rows)
        doc.add_paragraph()
    # pipes
    if "\u7ba1\u5f84\u9009\u578b" in collected:
        _add_heading(doc, "\u4e03\u3001\u7ba1\u5f84\u9009\u578b", level=1)
        rd = collected["\u7ba1\u5f84\u9009\u578b"]
        pipes = rd.get("pipes", [])
        headers = ["\u7ba1\u8def\u7c7b\u578b", "\u6d41\u91cf(m\u00b3/h)", "\u63a8\u8350DN", "\u5185\u5f84(mm)", "\u6d41\u901f(m/s)", "\u8bc4\u4ef7"]
        rows = [[p["pipe_type"], f"{p['flow']:.1f}", p["recommended_dn"], p["inner_diameter_mm"], f"{p['actual_velocity']:.2f}", p["advice"]] for p in pipes]
        _add_data_table(doc, headers, rows)
        doc.add_paragraph()
    # CIP
    if "CIP\u6e05\u6d17" in collected:
        _add_heading(doc, "\u516b\u3001CIP\u6e05\u6d17\u8ba1\u7b97", level=1)
        _add_kv_table(doc, collected["CIP\u6e05\u6d17"])
        doc.add_paragraph()
    # bubble
    if "\u6ce1\u70b9\u68c0\u6d4b" in collected:
        _add_heading(doc, "\u4e5d\u3001\u6ce1\u70b9\u68c0\u6d4b\u6c47\u603b", level=1)
        rd = collected["\u6ce1\u70b9\u68c0\u6d4b"]
        summary = {k: v for k, v in rd.items() if k != "records"}
        _add_kv_table(doc, summary)
        doc.add_paragraph()
    # BOM
    if "BOM\u6e05\u5355" in collected:
        _add_heading(doc, "\u5341\u3001BOM\u6e05\u5355", level=1)
        rd = collected["BOM\u6e05\u5355"]
        items = rd.get("items", [])
        headers = ["\u5e8f\u53f7", "\u540d\u79f0", "\u89c4\u683c", "\u6570\u91cf", "\u5355\u4f4d", "\u5355\u4ef7(\u5143)", "\u91d1\u989d(\u5143)", "\u5907\u6ce8"]
        rows = [[str(item.get(k, "")) for k in ["seq", "name", "spec", "quantity", "unit", "price", "amount", "remark"]] for item in items]
        _add_data_table(doc, headers, rows)
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f"\u603b\u4ef7: \uffe5{rd.get('total', 0):,.2f} \u5143")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)
